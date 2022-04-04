import email
import imaplib
import os
from email.header import decode_header
import yaml
import docx

import sql

with open("config.yaml") as file:
    config_list = list(yaml.load(file, Loader=yaml.FullLoader).values())

with open("config.yaml") as file:
    config_list = list(yaml.load(file, Loader=yaml.FullLoader).values())

for config_item in config_list[2:]:
    assert config_item, "Значение не может быть пустым"

LOG_SYMBOL, DEBUG, OUR_USERNAME, OUR_PASSWORD, SEND_USERNAME, SEND_PASSWORD, DATABASE_PATH = config_list[:7]
if not os.path.exists(DATABASE_PATH):
    print("БД не существует, будет создана новая БД")


def get_headers_email(email_numbers, read_mode="ALL"):
    # ALL: Выводит адрес, тему письма и дату
    # E_ONLY: Список адресов
    #

    mail = imaplib.IMAP4_SSL('imap.yandex.ru')
    if DEBUG:
        print(LOG_SYMBOL + "IMAP-подключение: успешно!")
    mail.login(OUR_USERNAME, OUR_PASSWORD)
    if DEBUG:
        print(LOG_SYMBOL + "IMAP-авторизация: успешно!")
    mail.list()
    mail.select("inbox")
    result, data = mail.search(None, "ALL")
    ids = data[0]
    id_list = ids.split()
    emails_list = id_list[len(id_list) - email_numbers:]
    headers_list = []
    for item in range(len(emails_list)):
        result, data = mail.fetch(emails_list[item], "(RFC822)")
        raw_email = data[0][1]
        raw_email_string = raw_email.decode('cp1251')
        email_message = email.message_from_string(raw_email_string)
        temp_list = {}
        if read_mode == "ALL":
            try:
                subject = email_message['Subject']
                byte_string, encoding = decode_header(subject)[0]
                subject = byte_string.decode(encoding)
            except TypeError:
                subject = ""
            temp_list = {"Email": email.utils.parseaddr(email_message['From'])[1],
                         "Date": email_message["Date"],
                         "Subject": subject,
                         "Id": data[0][0]}
        if read_mode == "E_ONLY":
            temp_list = email.utils.parseaddr(email_message['From'])[1]
        headers_list.append(temp_list)
    return headers_list


def read_work_mail():
    email_list = [sql.query_executor_select("SELECT Email FROM Clients")[i][0] for i in
                  range(len(sql.query_executor_select(
                      "SELECT Email FROM Clients")))]
    for item in get_headers_email(10, ):
        if item["Email"] in email_list:
            id = item["Id"].decode("utf-8").split(" ")[0]
            mail = imaplib.IMAP4_SSL('imap.yandex.ru')
            mail.login(OUR_USERNAME, OUR_PASSWORD)
            mail.list()
            mail.select("inbox")
            result, data = mail.fetch(id, "(RFC822)")
            workmail = email.message_from_bytes(data[0][1])
            filename = ""
            if workmail.is_multipart():
                for part in workmail.walk():
                    filename = part.get_filename()
                    if filename:
                        with open(part.get_filename(), 'wb') as new_file:
                            new_file.write(part.get_payload(decode=True))
            if item["Subject"] == "Заполненный документ " + "НВ":
                doc = docx.Document(filename)
                table = doc.tables[0]
                date = table.cell(0, 1).text
                purpose = table.cell(5, 1).text
                sql.query_executor_insert(
                    "UPDATE Meetings SET Purpose = \"{}\" WHERE Date = \"{}\"".format(purpose, date))
                os.remove(filename)
            if item["Subject"] == "Заполненный документ " + "НК":
                doc = docx.Document(filename)
                table = doc.tables[0]
                fullname = table.cell(0, 1).text + " " + table.cell(1, 1).text + " " + table.cell(2, 1).text
                phone = table.cell(3, 1).text
                ppp = table.cell(4, 1).text
                company_name = table.cell(5, 1).text
                sql.query_executor_insert(
                    "UPDATE Clients SET Name = \"{}\", PhoneNumber = \"{}\", Company = \"{}\" WHERE Email = \"{}\"".format(
                        fullname, phone, company_name, ppp))
                os.remove(filename)
            if item["Subject"] == "Заполненный документ " + "ЗС":
                doc = docx.Document(filename)
                table = doc.tables[0]
                fullname = table.cell(0, 2).text
                table = doc.tables[1]
                print(fullname)
                type_product = table.cell(0, 1).text
                value = table.cell(1, 1).text
                date_open = table.cell(2, 1).text
                date_close = table.cell(3, 1).text
                order_id = sql.query_executor_select(
                    "SELECT Order_Id FROM Orders WHERE Client_ID = (SELECT Client_ID From Clients WHERE Name = \"{"
                    "}\") AND Date_Close = \"\"".format(
                        fullname))[0][0]
                sql.query_executor_insert("UPDATE Orders SET Value = \"{}\", Type = \"{}\", Date_Open = \"{}\", "
                                          "Date_Close = \"{}\" WHERE Order_ID = {}".format(value, type_product,
                                                                                           date_open, date_close,
                                                                                           order_id))
                os.remove(filename)
                print("Done")


if __name__ == "__main__":
    read_work_mail()
