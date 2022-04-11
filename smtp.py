import datetime
import os
import sqlite3
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from smtplib import *
import docx
import yaml

import imap
import sql

with open("config.yaml") as file:
    config_list = list(yaml.load(file, Loader=yaml.FullLoader).values())

for config_item in config_list[2:]:
    assert config_item, "Значение не может быть пустым"

LOG_SYMBOL, DEBUG, OUR_USERNAME, OUR_PASSWORD, SEND_USERNAME, SEND_PASSWORD, DATABASE_PATH, DOCUMENTS_PATH = config_list
if not os.path.exists(DATABASE_PATH):
    print("БД не существует, будет создана новая БД")


def send_message_clients(to_email=SEND_USERNAME):
    server = SMTP("smtp.gmail.com", 587)
    if DEBUG:
        print(LOG_SYMBOL + "SMTP-подключение: подключение!")
    server.ehlo()
    server.starttls()
    msg = MIMEMultipart()
    server.login(SEND_USERNAME, SEND_PASSWORD)

    if DEBUG:
        print(LOG_SYMBOL + "SMTP-авторизация: подключение!")
    if sql.get_client_data(to_email) is None:
        print(LOG_SYMBOL + "Такого клиента в нашей БД нет")
        server.quit()
    else:
        name, lead_type = sql.get_client_data(to_email)
        if lead_type == "M":
            msg['Subject'] = "Знакомство с нами"
            msg['From'] = SEND_USERNAME
            msg['To'] = to_email
            body = open("emails_forms/email_form_m.txt", encoding="utf-8").read().format(name)
            msg.attach(MIMEText(body, 'plain'))
            attach = MIMEApplication(open(DOCUMENTS_PATH + "New client.docx", 'rb').read())
            attach.add_header('Content-Disposition', 'attachment', filename='New client.docx')
            msg.attach(attach)
            server.sendmail(SEND_USERNAME, to_email, msg.as_string())
            if DEBUG:
                print(LOG_SYMBOL + "Отправлено письмо клиенту категории {}".format(lead_type))
            server.quit()
        if lead_type == "L":

            providers_rows = \
                sql.query_executor_select("SELECT Name, Price FROM Providers WHERE Provider_ID NOT IN (SELECT "
                                          "Provider_ID FROM "
                                          "Orders Where Status = \"Open\") ORDER BY Priority")
            logists_rows = sql.query_executor_select("SELECT Name, Price FROM Logists WHERE Logist_ID NOT IN (SELECT "
                                                     "Logist_ID FROM "
                                                     "Orders Where Status = \"Open\") ORDER BY Priority")
            if not providers_rows:
                providers_rows = [("", 0)]
            if not logists_rows:
                logists_rows = [("", 0)]

            providers_rows = providers_rows[0]
            logists_rows = logists_rows[0]

            doc = docx.Document(DOCUMENTS_PATH + "Deal contract.docx")
            table = doc.tables[1]
            table.cell(2, 1).text = datetime.datetime.now().strftime("%Y-%m-%d")
            table.cell(4, 1).text = providers_rows[0]
            table.cell(5, 1).text = logists_rows[0]
            table.cell(6, 1).text = "Около {} у.e (при объеме до 1000 шт.)".format(
                str(logists_rows[1] * providers_rows[1] * 1.15))
            doc.save("temp_doc.docx")

            logist = sql.query_executor_select(
                "SELECT Logist_ID FROM Logists WHERE Name = \"{}\"".format(logists_rows[0]))
            if not logist:
                logist = ""
            else:
                logist = logist[0][0]

            provider = sql.query_executor_select(
                "SELECT Provider_ID FROM Providers WHERE Name = \"{}\"".format(providers_rows[0]))
            if not provider:
                provider = ""
            else:
                provider = provider[0][0]

            new_order = [sql.query_executor_select("SELECT Order_ID From Orders ORDER BY Order_ID DESC")[0][0] + 1,
                         int(sql.query_executor_select(
                             "SELECT Client_ID from Clients WHERE Email = \"{}\"".format("bunsanandme@yandex.ru"))[0][
                                 0]), "", "", datetime.datetime.now().strftime("%Y-%m-%d"), "", "Open",
                         logist,
                         provider]
            sql.query_executor_insert("INSERT INTO Orders VALUES {}".format(tuple(new_order)))

            msg['Subject'] = "Заключение сделки"
            msg['From'] = SEND_USERNAME
            msg['To'] = to_email
            body = open("emails_forms/email_form_l.txt", encoding="utf-8").read().format(name)
            msg.attach(MIMEText(body, 'plain'))
            attach = MIMEApplication(open("temp_doc.docx", 'rb').read())
            attach.add_header('Content-Disposition', 'attachment', filename='Заключение сделки.docx')
            msg.attach(attach)
            server.sendmail(SEND_USERNAME, to_email, msg.as_string())
            if DEBUG:
                print(LOG_SYMBOL + "Отправлено письмо клиенту категории {}".format(lead_type))
            os.remove("temp_doc.docx")
            server.quit()


def send_message_meetings(to_email=SEND_USERNAME,
                          meeting_date=datetime.datetime.now().strftime("%Y-%m-%d"),
                          address="Наш офис"):
    server = SMTP("smtp.gmail.com", 587)
    if DEBUG:
        print(LOG_SYMBOL + "SMTP-подключение: подключение!")
    server.ehlo()
    server.starttls()
    msg = MIMEMultipart()
    server.login(SEND_USERNAME, SEND_PASSWORD)
    if DEBUG:
        print(LOG_SYMBOL + "SMTP-авторизация: подключение!")
    if to_email not in imap.get_headers_email(5, "E_ONLY"):
        print(LOG_SYMBOL + "Такого клиента в нашей БД нет")
        server.quit()
    else:
        doc = docx.Document(DOCUMENTS_PATH + "Meeting.docx")
        table = doc.tables[0]
        table.cell(0, 1).text = meeting_date
        table.cell(1, 1).text = address
        doc.save("temp_doc.docx")
        new_meeting = [sql.query_executor_select("SELECT Meeting_ID From Meetings ORDER BY Meeting_ID DESC")[0][0] + 1,
                       sql.query_executor_select("SELECT Client_ID FROM Clients WHERE Email = \"{}\"".format(to_email))[
                           0][
                           0], meeting_date, "", address]
        try:
            sql.query_executor_insert("INSERT INTO Meetings VALUES {}".format(tuple(new_meeting)))
        except sqlite3.IntegrityError:
            print("На эту дату уже назначена встреча.")
            server.quit()
            return None
        name, lead_type = sql.get_client_data(to_email)
        msg['Subject'] = "Приглашение на встречу"
        msg['From'] = SEND_USERNAME
        msg['To'] = to_email
        body = "https://forms.gle/igmb9SSzKvW9TksQ9 "
        body += open("emails_forms/email_form_meeting.txt", encoding="utf-8").read().format(name)
        msg.attach(MIMEText(body, 'plain'))
        attach = MIMEApplication(open("temp_doc.docx", 'rb').read())
        attach.add_header('Content-Disposition', 'attachment', filename='Meeting.docx')
        msg.attach(attach)
        server.sendmail(SEND_USERNAME, to_email, msg.as_string())
        if DEBUG:
            print(LOG_SYMBOL + "Отправлено письмо клиенту категории {}".format(lead_type))
        os.remove("temp_doc.docx")
        server.quit()


if __name__ == "__main__":
    send_message_meetings("bunsanandme@yandex.ru")
