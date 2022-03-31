import sqlite3
import imaplib
import email
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from smtplib import *
import docx
import datetime
import os
from email.header import decode_header
import yaml
from pprint import pprint

LOG_SYMBOL = ">> "
DEBUG = False

with open("config.yaml") as file:
    config_list = list(yaml.load(file, Loader=yaml.FullLoader).values())

OUR_USERNAME, OUR_PASSWORD, SEND_USERNAME, SEND_PASSWORD, DATABASE_PATH, DOCUMENTS_PATH = config_list


# -------
# SQL
# -------

def query_executor_select(query):
    sqlite_connection = None
    try:
        sqlite_connection = sqlite3.connect(DATABASE_PATH)
        cursor = sqlite_connection.cursor()
        cursor.execute(query)
        rows = cursor.fetchall()
        cursor.close()
        return rows
    except sqlite3.Error as error:
        if DEBUG:
            print(LOG_SYMBOL + "Ошибка при подключении к sqlite")
            print(LOG_SYMBOL + "Ошибка:", error)
    finally:
        if sqlite_connection:
            sqlite_connection.close()


def query_executor_insert(query):
    sqlite_connection = sqlite3.connect(DATABASE_PATH)
    cursor = sqlite_connection.cursor()
    cursor.execute(query)
    sqlite_connection.commit()
    return cursor.lastrowid


def get_client_data(client_email):
    sqlite_connection = None
    try:
        sqlite_connection = sqlite3.connect(DATABASE_PATH)
        cursor = sqlite_connection.cursor()
        if DEBUG:
            print(LOG_SYMBOL + "Подключение прошло успешно!")
        cursor.execute("SELECT * FROM Clients WHERE Email = \"{}\"".format(client_email))
        if DEBUG:
            print(LOG_SYMBOL + "Выполнен запрос: успешно!")
        try:
            rows = cursor.fetchall()[0]
        except IndexError:
            print(LOG_SYMBOL + "ВНИМАНИЕ: Запрос пустой!")
            return None
        cursor.close()
        return rows[3], rows[7]

    except sqlite3.Error as error:
        if DEBUG:
            print(LOG_SYMBOL + "Ошибка при подключении к sqlite")
            print(LOG_SYMBOL + "Ошибка:", error)
    finally:
        if sqlite_connection:
            sqlite_connection.close()
            if DEBUG:
                print(LOG_SYMBOL + "Отключение от БД...")


# -------
# IMAP
# -------

def get_headers_email(email_numbers, read_mode="ALL"):
    # ALL: Выводит адрес, тему письма и дату
    # E_ONLY: Список адресов
    #

    mail = imaplib.IMAP4_SSL('imap.yandex.ru')
    if DEBUG:
        print(LOG_SYMBOL + "IMAP-подключение: успешно!")
    mail.login(OUR_USERNAME, OUR_PASSWORD)
    if DEBUG:
        print(LOG_SYMBOL + "IMAP-авторизация: успешно!")
    mail.list()
    mail.select("inbox")
    result, data = mail.search(None, "ALL")
    ids = data[0]
    id_list = ids.split()
    emails_list = id_list[len(id_list) - email_numbers:]
    headers_list = []
    for item in range(len(emails_list)):
        result, data = mail.fetch(emails_list[item], "(RFC822)")
        raw_email = data[0][1]
        raw_email_string = raw_email.decode('cp1251')
        email_message = email.message_from_string(raw_email_string)
        temp_list = {}
        if read_mode == "ALL":
            try:
                subject = email_message['Subject']
                byte_string, encoding = decode_header(subject)[0]
                subject = byte_string.decode(encoding)
            except TypeError:
                subject = ""

            temp_list = {"Email": email.utils.parseaddr(email_message['From'])[1],
                         "Date": email_message["Date"],
                         "Subject": subject,
                         "Id": data[0][0]}
        if read_mode == "E_ONLY":
            temp_list = email.utils.parseaddr(email_message['From'])[1]
        headers_list.append(temp_list)
    return headers_list


def read_work_mail():
    email_list = [query_executor_select("SELECT Email FROM Clients")[i][0] for i in range(len(query_executor_select(
        "SELECT Email FROM Clients")))]
    for item in get_headers_email(10, ):
        if item["Email"] in email_list:
            id = item["Id"].decode("utf-8").split(" ")[0]
            mail = imaplib.IMAP4_SSL('imap.yandex.ru')
            mail.login(OUR_USERNAME, OUR_PASSWORD)
            mail.list()
            mail.select("inbox")
            result, data = mail.fetch(id, "(RFC822)")
            workmail = email.message_from_bytes(data[0][1])
            filename = ""
            if workmail.is_multipart():
                for part in workmail.walk():
                    content_type = part.get_content_type()
                    filename = part.get_filename()
                    if filename:
                        with open(part.get_filename(), 'wb') as new_file:
                            new_file.write(part.get_payload(decode=True))
            if item["Subject"] == "Заполненный документ " + "НВ":
                doc = docx.Document(filename)
                table = doc.tables[0]
                date = table.cell(0, 1).text
                purpose = table.cell(5, 1).text
                query_executor_insert("UPDATE Meetings SET Purpose = \"{}\" WHERE Date = \"{}\"".format(purpose, date))
                os.remove(filename)
            if item["Subject"] == "Заполненный документ " + "НК":
                doc = docx.Document(filename)
                table = doc.tables[0]
                fullname = table.cell(0, 1).text + " " + table.cell(1, 1).text + " " + table.cell(2, 1).text
                phone = table.cell(3, 1).text
                ppp = table.cell(4, 1).text
                company_name = table.cell(5, 1).text
                query_executor_insert(
                    "UPDATE Clients SET Name = \"{}\", PhoneNumber = \"{}\", Company = \"{}\" WHERE Email = \"{}\"".format(
                        fullname, phone, company_name, ppp))
                os.remove(filename)
            if item["Subject"] == "Заполненный документ " + "ЗС":
                doc = docx.Document(filename)
                table = doc.tables[0]
                fullname = table.cell(0, 2).text
                table = doc.tables[1]
                type_product = table.cell(0, 1).text
                value = table.cell(1, 1).text
                date_open = table.cell(2, 1).text
                date_close = table.cell(3, 1).text
                order_id = query_executor_select(
                    "SELECT Order_Id FROM Orders WHERE Client_ID = (SELECT Client_ID From Clients WHERE Name = \"{"
                    "}\") AND Date_Close = \"\"".format(
                        fullname))[0][0]
                query_executor_insert("UPDATE Orders SET Value = \"{}\", Type = \"{}\", Date_Open = \"{}\", "
                                      "Date_Close = \"{}\" WHERE Order_ID = {}".format(value, type_product,
                                                                                       date_open, date_close,
                                                                                       order_id))
                os.remove(filename)


# -------
# SMTP
# -------

def send_message_clients(
        to_email=SEND_USERNAME):
    server = SMTP("smtp.gmail.com", 587)
    if DEBUG:
        print(LOG_SYMBOL + "SMTP-подключение: подключение!")
    server.ehlo()
    server.starttls()
    msg = MIMEMultipart()
    server.login(SEND_USERNAME, SEND_PASSWORD)

    if DEBUG:
        print(LOG_SYMBOL + "SMTP-авторизация: подключение!")
    if to_email not in get_headers_email(5, "E_ONLY"):
        print(LOG_SYMBOL + "Такого клиента в нашей БД нет")
        server.quit()
    else:
        name, lead_type = get_client_data(to_email)
        if lead_type == "M":
            msg['Subject'] = "Знакомство с нами"
            msg['From'] = SEND_USERNAME
            msg['To'] = to_email
            body = "Здравствуйте, {}!" \
                   "\nРады, что вы решили обратиться за помощью к нашей компании. Давайте познакомимся!\n" \
                   "Заполните документ ниже и пришлите ответом с темой \"Заполненый документ НК\". Пожалуйста, " \
                   "не меняйте название файла - боимся вас потерять\n\n\n" \
                   "----\n" \
                   "Компания\n" \
                   "Телефон".format(name, )
            msg.attach(MIMEText(body, 'plain'))
            attach = MIMEApplication(open("New client.docx", 'rb').read())
            attach.add_header('Content-Disposition', 'attachment', filename='New client.docx')
            msg.attach(attach)
            server.sendmail(SEND_USERNAME, to_email, msg.as_string())
            if DEBUG:
                print(LOG_SYMBOL + "Отправлено письмо клиенту категории {}".format(lead_type))

            server.quit()
        if lead_type == "L":
            providers_rows = query_executor_select("SELECT Name, Price FROM Providers WHERE Provider_ID NOT IN (SELECT "
                                                   "Provider_ID FROM "
                                                   "Orders Where Status = \"Open\") ORDER BY Priority")[0]
            logists_rows = query_executor_select("SELECT Name, Price FROM Logists WHERE Logist_ID NOT IN (SELECT "
                                                 "Logist_ID FROM "
                                                 "Orders Where Status = \"Open\") ORDER BY Priority")[0]
            doc = docx.Document("T:\\my\\WORK\\automation\\automation\\Deal contract.docx")
            table = doc.tables[1]
            table.cell(2, 1).text = datetime.datetime.now().strftime("%Y-%m-%d")
            table.cell(4, 1).text = providers_rows[0]
            table.cell(5, 1).text = logists_rows[0]
            table.cell(6, 1).text = "Около {} у.e (при объеме до 1000 шт.)".format(
                str(logists_rows[1] * providers_rows[1] * 1.15))
            doc.save("temp_doc.docx")

            new_order = [query_executor_select("SELECT Order_ID From Orders ORDER BY Order_ID DESC")[0][0] + 1,
                         int(query_executor_select(
                             "SELECT Client_ID from Clients WHERE Email = \"{}\"".format("bunsanandme@yandex.ru"))[0][
                                 0]), "", "", datetime.datetime.now().strftime("%Y-%m-%d"), "", "Open",
                         query_executor_select(
                             "SELECT Logist_ID FROM Logists WHERE Name = \"{}\"".format(logists_rows[0]))[0][0],
                         query_executor_select(
                             "SELECT Provider_ID FROM Providers WHERE Name = \"{}\"".format(providers_rows[0]))[0][0]]
            query_executor_insert("INSERT INTO Orders VALUES {}".format(tuple(new_order)))

            msg['Subject'] = "Заключение сделки"
            msg['From'] = SEND_USERNAME
            msg['To'] = to_email
            body = "Здравствуйте, {}!" \
                   "\nПрикладываю к письму договор-форму о заключении сделки. Рассмотрите его, при возникновении " \
                   "вопросов обращайтесь.\n" \
                   "Заполните документ ниже и пришлите ответом с темой \"Заполненый документ ЗС\". Пожалуйста, " \
                   "не меняйте название файла - боимся вас потерять\n\n\n" \
                   "----\n" \
                   "Компания\n" \
                   "Телефон".format(name, )
            msg.attach(MIMEText(body, 'plain'))
            attach = MIMEApplication(open("temp_doc.docx", 'rb').read())
            attach.add_header('Content-Disposition', 'attachment', filename='Заключение сделки.docx')
            msg.attach(attach)
            server.sendmail(SEND_USERNAME, to_email, msg.as_string())
            if DEBUG:
                print(LOG_SYMBOL + "Отправлено письмо клиенту категории {}".format(lead_type))
            os.remove("temp_doc.docx")
            server.quit()


def send_message_meetings(to_email=SEND_USERNAME,
                          meeting_date=datetime.datetime.now().strftime("%Y-%m-%d"),
                          address="Наш офис"):
    server = SMTP("smtp.gmail.com", 587)
    if DEBUG:
        print(LOG_SYMBOL + "SMTP-подключение: подключение!")
    server.ehlo()
    server.starttls()
    msg = MIMEMultipart()
    server.login(SEND_USERNAME, SEND_PASSWORD)
    if DEBUG:
        print(LOG_SYMBOL + "SMTP-авторизация: подключение!")
    if to_email not in get_headers_email(5, "E_ONLY"):
        print(LOG_SYMBOL + "Такого клиента в нашей БД нет")
        server.quit()
    else:
        doc = docx.Document("T:\\my\\WORK\\automation\\automation\\Meeting.docx")
        table = doc.tables[0]
        table.cell(0, 1).text = meeting_date
        table.cell(1, 1).text = address
        doc.save("temp_doc.docx")
        new_meeting = [query_executor_select("SELECT Meeting_ID From Meetings ORDER BY Meeting_ID DESC")[0][0] + 1,
                       query_executor_select("SELECT Client_ID FROM Clients WHERE Email = \"{}\"".format(to_email))[0][
                           0], meeting_date, "", address]
        query_executor_insert("INSERT INTO Meetings VALUES {}".format(tuple(new_meeting)))
        name, lead_type = get_client_data(to_email)
        msg['Subject'] = "Приглашение на встречу"
        msg['From'] = SEND_USERNAME
        msg['To'] = to_email
        body = "Здравствуйте, {}!" \
               "\nПриглашаем вас на деловую встречу в наш офис по адресу.\n" \
               "Заполните документ ниже и пришлите ответом с темой \"Заполненый документ НВ\". Пожалуйста, " \
               "не меняйте название файла - боимся вас потерять\n\n\n" \
               "----\n" \
               "Компания\n" \
               "Телефон".format(name, )
        msg.attach(MIMEText(body, 'plain'))
        attach = MIMEApplication(open("temp_doc.docx", 'rb').read())
        attach.add_header('Content-Disposition', 'attachment', filename='Meeting.docx')
        msg.attach(attach)
        server.sendmail(SEND_USERNAME, to_email, msg.as_string())
        if DEBUG:
            print(LOG_SYMBOL + "Отправлено письмо клиенту категории {}".format(lead_type))
            os.remove("temp_doc.docx")
        server.quit()


if __name__ == "__main__":
    read_work_mail()
